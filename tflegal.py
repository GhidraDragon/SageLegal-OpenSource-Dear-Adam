#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import re
import pickle
import os
import datetime
import sqlite3
from collections import OrderedDict
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sklearn.feature_extraction.text import CountVectorizer
import numpy as np
import tensorflow as tf
from tensorflow.keras import layers, models

def read_input_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        from PyPDF2 import PdfReader
        with open(filepath, 'rb') as f:
            reader = PdfReader(f)
            lines = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    lines.extend(text.splitlines())
            return "\n".join(lines)
    elif ext == '.docx':
        from docx import Document as DocxDocument
        doc_loaded = DocxDocument(filepath)
        lines = []
        for para in doc_loaded.paragraphs:
            lines.append(para.text)
        return "\n".join(lines)
    elif ext == '.odt':
        from odf.opendocument import load
        from odf import text as odftext
        from odf import teletype
        odt_doc = load(filepath)
        lines = []
        for elem in odt_doc.getElementsByType(odftext.P):
            lines.append(teletype.extractText(elem))
        return "\n".join(lines)
    else:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()

def generate_smart_filename(original_filename, text, dt_string):
    vectorizer = CountVectorizer(stop_words='english', max_features=50)
    X = vectorizer.fit_transform([text])
    word_counts = X.toarray().sum(axis=0)
    sorted_indices = np.argsort(-word_counts)
    feature_names = vectorizer.get_feature_names_out()
    top_words = []
    for idx in sorted_indices[:3]:
        top_words.append(feature_names[idx])
    base, ext = os.path.splitext(original_filename)
    top_part = "_".join(top_words)
    return f"{base}_{top_part}_{dt_string}{ext}"

def is_exhibit_reference(line_str):
    return bool(re.search(r'\bEXHIBIT\s+\d+:', line_str, re.IGNORECASE))

def is_line_all_caps(line_str):
    if not re.search(r'[A-Z]', line_str):
        return False
    return not re.search(r'[a-z]', line_str)

def is_line_of_equals(line_str):
    s = line_str.strip()
    if len(s) < 5:
        return False
    return bool(re.match(r'^[=]+$', s))

def is_line_of_dashes(line_str):
    s = line_str.strip()
    if len(s) < 5:
        return False
    return bool(re.match(r'^-+$', s))

def detect_legal_title_blocks(lines):
    i = 0
    n = len(lines)
    while i < n:
        if is_line_of_equals(lines[i]):
            j = i + 1
            inner_lines = []
            found_bottom = False
            while j < n:
                if is_line_of_equals(lines[j]):
                    found_bottom = True
                    j += 1
                    break
                else:
                    inner_lines.append(lines[j])
                j += 1
            if found_bottom:
                yield ("legal_page_title_block", inner_lines)
                i = j
            else:
                yield ("delimiter_line", lines[i])
                i += 1
        elif is_line_of_dashes(lines[i]):
            yield ("delimiter_line", lines[i])
            i += 1
        else:
            yield ("normal_line", lines[i])
            i += 1

def wrap_text_to_lines(pdf_canvas, full_text, font_name, font_size, max_width):
    pdf_canvas.setFont(font_name, font_size)
    paragraphs = full_text.split('\n')
    all_lines = []
    for paragraph in paragraphs:
        words = paragraph.split()
        if not words:
            all_lines.append(("", False))
            continue
        current_line = ""
        for word in words:
            test_line = word if not current_line else (current_line + " " + word)
            if pdf_canvas.stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                all_lines.append((current_line, True))
                current_line = word
        if current_line:
            all_lines.append((current_line, False))
    return all_lines

def draw_firm_name_vertical_center(pdf_canvas, text, page_width, page_height):
    pdf_canvas.saveState()
    pdf_canvas.setFont("Helvetica-Bold", 10)
    text_width = pdf_canvas.stringWidth(text, "Helvetica-Bold", 10)
    x_pos = 0.2 * inch
    y_center = page_height / 2.0
    y_pos = y_center - (text_width / 2.0)
    pdf_canvas.translate(x_pos, y_pos)
    pdf_canvas.rotate(90)
    pdf_canvas.drawString(0, 0, text)
    pdf_canvas.restoreState()

def draw_legal_page_title_block(
    pdf_canvas,
    page_width,
    page_height,
    block_lines,
    firm_name,
    case_name,
    page_number,
    total_pages
):
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.3 * inch)
    draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
    pdf_canvas.setFont("Helvetica-Bold", 12)
    pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
    pdf_canvas.setLineWidth(1)
    pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
    pdf_canvas.setFont("Helvetica-Bold", 14)
    line_spacing = 0.3 * inch
    y_text = page_height - 1.5 * inch
    for line_str in block_lines:
        pdf_canvas.drawCentredString(page_width / 2.0, y_text, line_str)
        y_text -= line_spacing
    pdf_canvas.setFont("Helvetica-Oblique", 9)
    footer_text = f"Page {page_number} of {total_pages}"
    pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)

def draw_page_of_segments(
    pdf_canvas,
    page_width,
    page_height,
    segments,
    start_index,
    max_lines_per_page,
    firm_name,
    case_name,
    page_number,
    total_pages,
    line_offset_x,
    line_offset_y,
    line_spacing,
    heading_positions
):
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.3 * inch)
    draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
    pdf_canvas.setFont("Helvetica-Bold", 12)
    pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
    pdf_canvas.setLineWidth(1)
    pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
    end_index = start_index
    current_line_count = 0
    y_text = line_offset_y
    while end_index < len(segments) and current_line_count < max_lines_per_page:
        seg = segments[end_index]
        if seg.get("page_always_new"):
            if current_line_count > 0:
                break
            else:
                block_lines = seg["lines"]
                draw_legal_page_title_block(
                    pdf_canvas,
                    page_width,
                    page_height,
                    block_lines,
                    firm_name,
                    case_name,
                    page_number,
                    total_pages
                )
                end_index += 1
                return end_index
        line_number = end_index + 1
        pdf_canvas.setFont("Helvetica", 10)
        pdf_canvas.drawString(line_offset_x - 0.6 * inch, y_text, str(line_number))
        pdf_canvas.drawString(page_width - 0.4 * inch, y_text, str(line_number))
        pdf_canvas.setFont(seg["font_name"], seg["font_size"])
        if seg["is_heading"] or seg["is_subheading"]:
            heading_positions.append(
                (
                    seg["text"],
                    page_number,
                    line_number,
                    seg["is_subheading"]
                )
            )
        if seg.get("delimiter_line"):
            pdf_canvas.setLineWidth(1)
            pdf_canvas.line(line_offset_x, y_text + 4, page_width - 0.5 * inch, y_text + 4)
            y_text -= line_spacing
            current_line_count += 1
            end_index += 1
            continue
        if seg["alignment"] == "center":
            left_boundary = line_offset_x
            right_boundary = page_width - 0.5 * inch
            mid_x = (left_boundary + right_boundary) / 2.0
            pdf_canvas.drawCentredString(mid_x, y_text, seg["text"])
        else:
            pdf_canvas.drawString(line_offset_x, y_text, seg["text"])
        y_text -= line_spacing
        current_line_count += 1
        end_index += 1
    pdf_canvas.setFont("Helvetica-Oblique", 9)
    footer_text = f"Page {page_number} of {total_pages}"
    pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)
    return end_index

def generate_index_pdf(index_filename, firm_name, case_name, heading_positions):
    pdf_canvas = canvas.Canvas(index_filename, pagesize=letter)
    pdf_canvas.setTitle("Table of Contents")
    page_width, page_height = letter
    top_margin = 1.0 * inch
    bottom_margin = 1.0 * inch
    left_margin = 1.0 * inch
    right_margin = 0.5 * inch
    line_spacing = 0.25 * inch
    temp_c = canvas.Canvas("dummy.pdf", pagesize=letter)

    def wrap_text(linestr, font_name, font_size, maxwidth):
        temp_c.setFont(font_name, font_size)
        return wrap_text_to_lines(temp_c, linestr, font_name, font_size, maxwidth)

    max_entry_width = page_width - left_margin - 1.5 * inch
    flattened_lines = []
    for (heading_text, pg_num, ln_num, is_sub) in heading_positions:
        if is_sub:
            font_name = "Helvetica"
            font_size = 9
        else:
            font_name = "Helvetica-Bold"
            font_size = 10
        wrapped = wrap_text(heading_text, font_name, font_size, max_entry_width)
        text_lines = [w[0] for w in wrapped] if wrapped else [""]
        for i, txt_line in enumerate(text_lines):
            flattened_lines.append(
                (
                    txt_line,
                    pg_num,
                    ln_num,
                    font_name,
                    font_size,
                    (i == 0)
                )
            )
    usable_height = page_height - (top_margin + bottom_margin) - 1.0 * inch
    max_lines_per_page = int(usable_height // line_spacing)
    total_lines = len(flattened_lines)
    total_index_pages = max(1, (total_lines + max_lines_per_page - 1) // max_lines_per_page)
    i = 0
    current_page_index = 1
    while i < total_lines:
        pdf_canvas.setLineWidth(2)
        pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.3 * inch)
        draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
        pdf_canvas.setFont("Helvetica-Bold", 12)
        pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
        pdf_canvas.setLineWidth(1)
        pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
        pdf_canvas.setFont("Helvetica-Bold", 14)
        pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.75 * inch, "TABLE OF CONTENTS")
        x_text = 1.0 * inch
        y_text = page_height - top_margin - 0.75 * inch
        lines_on_this_page = 0
        while i < total_lines and lines_on_this_page < max_lines_per_page:
            line_text, pg_num, ln_num, font_name, font_size, show_pageline = flattened_lines[i]
            pdf_canvas.setFont(font_name, font_size)
            pdf_canvas.drawString(x_text, y_text, line_text)
            if show_pageline:
                label_str = f"{pg_num}:{ln_num}"
                pdf_canvas.drawRightString(page_width - 0.5 * inch - 0.2 * inch, y_text, label_str)
            y_text -= line_spacing
            i += 1
            lines_on_this_page += 1
        pdf_canvas.setFont("Helvetica-Oblique", 9)
        footer_text = f"Index Page {current_page_index} of {total_index_pages}"
        pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)
        if i < total_lines:
            pdf_canvas.showPage()
            current_page_index += 1
        else:
            break
    pdf_canvas.save()

def generate_complaint_docx(docx_filename, firm_name, case_name, header_od, sections_od, heading_styles):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    top_par = doc.add_paragraph()
    top_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top_par.add_run(f"{firm_name} | {case_name}\n")
    run.bold = True
    run.font.size = Pt(14)

    header_content = header_od.get("content", "")
    header_lines = header_content.splitlines()
    buffer_of_lines = []
    for kind, block_lines in detect_legal_title_blocks(header_lines):
        if kind == "legal_page_title_block":
            if buffer_of_lines:
                for line in buffer_of_lines:
                    p = doc.add_paragraph()
                    line_stripped = line.strip()
                    r = p.add_run(line_stripped)
                    if is_line_all_caps(line_stripped):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if is_exhibit_reference(line_stripped):
                        r.bold = True
                buffer_of_lines = []
            for line in block_lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                line_stripped = line.strip()
                runx = p.add_run(line_stripped)
                runx.bold = True
                runx.font.size = Pt(14)
                if is_exhibit_reference(line_stripped):
                    runx.bold = True
        elif kind == "delimiter_line":
            if buffer_of_lines:
                for line in buffer_of_lines:
                    p = doc.add_paragraph()
                    line_stripped = line.strip()
                    r = p.add_run(line_stripped)
                    if is_line_all_caps(line_stripped):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if is_exhibit_reference(line_stripped):
                        r.bold = True
                buffer_of_lines = []
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("――――――――――――――――――――――――――――――――――")
            r.bold = False
        else:
            buffer_of_lines.append(block_lines)

    if buffer_of_lines:
        for line in buffer_of_lines:
            p = doc.add_paragraph()
            line_stripped = line.strip()
            r = p.add_run(line_stripped)
            if is_line_all_caps(line_stripped):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if is_exhibit_reference(line_stripped):
                r.bold = True
        buffer_of_lines = []

    for section_key, section_body in sections_od.items():
        style_type = heading_styles.get(section_key, "section")
        doc.add_paragraph()
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if style_type == "section":
            run = heading_para.add_run(section_key)
            run.bold = True
            run.font.size = Pt(12)
        else:
            run = heading_para.add_run(section_key)
            run.bold = False
            run.font.size = Pt(11)
        if is_exhibit_reference(section_key):
            run.bold = True
        body_lines = section_body.splitlines()
        normal_buffer = []
        for kind, block_lines in detect_legal_title_blocks(body_lines):
            if kind == "legal_page_title_block":
                if normal_buffer:
                    for bline in normal_buffer:
                        bline_str = bline.strip()
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        rr = p.add_run(bline_str)
                        if style_type == "section":
                            rr.font.size = Pt(12)
                        else:
                            rr.font.size = Pt(11)
                        if is_exhibit_reference(bline_str):
                            rr.bold = True
                    normal_buffer = []
                for xline in block_lines:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    xline_str = xline.strip()
                    runx = p.add_run(xline_str)
                    runx.bold = True
                    runx.font.size = Pt(14)
                    if is_exhibit_reference(xline_str):
                        runx.bold = True
            elif kind == "delimiter_line":
                if normal_buffer:
                    for bline in normal_buffer:
                        bline_str = bline.strip()
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        rr = p.add_run(bline_str)
                        if style_type == "section":
                            rr.font.size = Pt(12)
                        else:
                            rr.font.size = Pt(11)
                        if is_exhibit_reference(bline_str):
                            rr.bold = True
                    normal_buffer = []
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = p.add_run("――――――――――――――――――――――――――――――――――")
                rr.bold = False
            else:
                normal_buffer.append(block_lines)
        if normal_buffer:
            for bline in normal_buffer:
                bline_str = bline.strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                rr = p.add_run(bline_str)
                if style_type == "section":
                    rr.font.size = Pt(12)
                else:
                    rr.font.size = Pt(11)
                if is_exhibit_reference(bline_str):
                    rr.bold = True
            normal_buffer = []
    doc.save(docx_filename)

def generate_toc_docx(docx_filename, firm_name, case_name, heading_positions):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    top_par = doc.add_paragraph()
    top_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top_par.add_run(f"{firm_name} | {case_name}\nTABLE OF CONTENTS\n")
    run.bold = True
    run.font.size = Pt(14)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for (heading_text, pg_num, ln_num, is_sub) in heading_positions:
        if re.match(r'(?i)^EXHIBIT\s+\d+:', heading_text):
            row_cells = table.add_row().cells
            left_cell = row_cells[0]
            right_cell = row_cells[1]
            left_par = left_cell.paragraphs[0]
            run_left = left_par.add_run(heading_text)
            run_left.font.size = Pt(12)
            run_left.bold = True
            right_par = right_cell.paragraphs[0]
            run_right = right_par.add_run(f"{pg_num}")
            run_right.font.size = Pt(12)
            run_right.bold = False
            right_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            row_cells = table.add_row().cells
            left_cell = row_cells[0]
            right_cell = row_cells[1]
            if is_sub:
                this_font_size = 11
                this_bold = False
            else:
                this_font_size = 12
                this_bold = True
            left_par = left_cell.paragraphs[0]
            run_left = left_par.add_run(heading_text)
            run_left.font.size = Pt(this_font_size)
            run_left.bold = this_bold
            if is_exhibit_reference(heading_text):
                run_left.bold = True
            right_par = right_cell.paragraphs[0]
            run_right = right_par.add_run(f"{pg_num}:{ln_num}")
            run_right.font.size = Pt(this_font_size)
            run_right.bold = False
            right_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(docx_filename)

def parse_documents_from_text(raw_text):
    lines = raw_text.splitlines()
    docs = []
    i = 0
    n = len(lines)
    while i < n:
        if is_line_of_equals(lines[i]) or is_line_of_dashes(lines[i]):
            j = i + 1
            doc_lines = []
            while j < n and not (is_line_of_equals(lines[j]) or is_line_of_dashes(lines[j])):
                doc_lines.append(lines[j])
                j += 1
            if j < n:
                docs.append("\n".join(doc_lines))
                i = j + 1
            else:
                break
        else:
            i += 1
    return docs

def parse_header_and_sections(raw_text):
    header_od = OrderedDict()
    sections_od = OrderedDict()
    heading_pattern = re.compile(r'^((?:[IVXLCDM]+\.|[0-9]+\.)+)\s+(.*)$', re.IGNORECASE)
    lines = raw_text.splitlines()
    idx = 0
    header_lines = []
    while idx < len(lines):
        line = lines[idx].rstrip('\n').rstrip('\r')
        if heading_pattern.match(line):
            break
        elif is_line_all_caps(line.strip()):
            break
        elif re.match(r'^[0-9]+\.\s*$', line.strip()):
            break
        header_lines.append(line)
        idx += 1
    header_od["content"] = "\n".join(header_lines)
    current_heading_key = None
    current_body_lines = []
    while idx < len(lines):
        line = lines[idx].rstrip('\n').rstrip('\r')
        m = heading_pattern.match(line)
        if m:
            if current_heading_key is not None:
                sections_od[current_heading_key] = "\n".join(current_body_lines)
            current_body_lines = []
            heading_number = m.group(1).strip()
            heading_title = m.group(2).strip()
            if heading_number.endswith('.'):
                heading_number = heading_number[:-1]
            current_heading_key = f"{heading_number} {heading_title}"
        elif is_line_all_caps(line.strip()):
            if current_heading_key is not None:
                sections_od[current_heading_key] = "\n".join(current_body_lines)
            current_body_lines = []
            current_heading_key = line.strip()
        elif re.match(r'^[0-9]+\.\s*$', line.strip()):
            if current_heading_key is not None:
                sections_od[current_heading_key] = "\n".join(current_body_lines)
            current_body_lines = []
            current_heading_key = line.strip()
        else:
            current_body_lines.append(line)
        idx += 1
    if current_heading_key is not None:
        sections_od[current_heading_key] = "\n".join(current_body_lines)
    return header_od, sections_od

def classify_headings(sections_od):
    heading_styles = {}
    for full_key in sections_od.keys():
        parts = full_key.split(None, 1)
        if len(parts) == 2:
            heading_number, _heading_text = parts[0], parts[1]
        else:
            heading_number = parts[0]
        dot_count = heading_number.count('.')
        if dot_count > 1:
            heading_styles[full_key] = "subsection"
        else:
            heading_styles[full_key] = "section"
    return heading_styles

def prepare_main_pdf_segments(header_text, sections_od, heading_styles, pdf_canvas, max_text_width):
    segments = []
    header_lines = header_text.splitlines()
    normal_buffer = []

    def flush_normal_buffer():
        for line in normal_buffer:
            line_str = line.strip()
            if not line_str:
                segments.append({
                    "text": "",
                    "font_name": "Helvetica",
                    "font_size": 10,
                    "alignment": "left",
                    "is_heading": False,
                    "is_subheading": False
                })
            elif is_line_all_caps(line_str):
                wrapped = wrap_text_to_lines(pdf_canvas, line_str, "Helvetica", 10, max_text_width)
                for (wl, _) in wrapped:
                    if is_exhibit_reference(line_str):
                        segments.append({
                            "text": wl,
                            "font_name": "Helvetica-Bold",
                            "font_size": 10,
                            "alignment": "center",
                            "is_heading": False,
                            "is_subheading": False
                        })
                    else:
                        segments.append({
                            "text": wl,
                            "font_name": "Helvetica",
                            "font_size": 10,
                            "alignment": "center",
                            "is_heading": False,
                            "is_subheading": False
                        })
            else:
                wrapped = wrap_text_to_lines(pdf_canvas, line_str, "Helvetica", 10, max_text_width)
                for (wl, _) in wrapped:
                    if is_exhibit_reference(line_str):
                        segments.append({
                            "text": wl,
                            "font_name": "Helvetica-Bold",
                            "font_size": 10,
                            "alignment": "left",
                            "is_heading": False,
                            "is_subheading": False
                        })
                    else:
                        segments.append({
                            "text": wl,
                            "font_name": "Helvetica",
                            "font_size": 10,
                            "alignment": "left",
                            "is_heading": False,
                            "is_subheading": False
                        })
        normal_buffer.clear()

    for kind, block_lines in detect_legal_title_blocks(header_lines):
        if kind == "legal_page_title_block":
            flush_normal_buffer()
            lines_cleaned = [ln.strip() for ln in block_lines]
            segments.append({
                "legal_page_title": True,
                "page_always_new": True,
                "lines": lines_cleaned
            })
        elif kind == "delimiter_line":
            flush_normal_buffer()
            segments.append({
                "delimiter_line": True,
                "font_name": "Helvetica",
                "font_size": 10,
                "is_heading": False,
                "is_subheading": False
            })
        else:
            normal_buffer.append(block_lines)
    flush_normal_buffer()

    for section_key, section_body in sections_od.items():
        style = heading_styles.get(section_key, "section")
        if style == "section":
            heading_font_name = "Helvetica-Bold"
            heading_font_size = 10
            body_font_name = "Helvetica"
            body_font_size = 10
            is_heading = True
            is_subheading = False
        else:
            heading_font_name = "Helvetica"
            heading_font_size = 9
            body_font_name = "Helvetica"
            body_font_size = 9
            is_heading = False
            is_subheading = True

        segments.append({
            "text": "",
            "font_name": body_font_name,
            "font_size": body_font_size,
            "alignment": "left",
            "is_heading": False,
            "is_subheading": False
        })
        if is_exhibit_reference(section_key):
            heading_wrapped = wrap_text_to_lines(pdf_canvas, section_key, "Helvetica-Bold", heading_font_size, max_text_width)
            for (wl, _) in heading_wrapped:
                segments.append({
                    "text": wl,
                    "font_name": "Helvetica-Bold",
                    "font_size": heading_font_size,
                    "alignment": "center",
                    "is_heading": is_heading,
                    "is_subheading": is_subheading
                })
        else:
            heading_wrapped = wrap_text_to_lines(pdf_canvas, section_key, heading_font_name, heading_font_size, max_text_width)
            for (wl, _) in heading_wrapped:
                segments.append({
                    "text": wl,
                    "font_name": heading_font_name,
                    "font_size": heading_font_size,
                    "alignment": "center",
                    "is_heading": is_heading,
                    "is_subheading": is_subheading
                })
        lines_of_body = section_body.splitlines()
        normal_buffer_sec = []

        def flush_section_buffer():
            for line in normal_buffer_sec:
                line_str = line.strip()
                if not line_str:
                    segments.append({
                        "text": "",
                        "font_name": body_font_name,
                        "font_size": body_font_size,
                        "alignment": "left",
                        "is_heading": False,
                        "is_subheading": False
                    })
                else:
                    wrapped = wrap_text_to_lines(pdf_canvas, line_str, body_font_name, body_font_size, max_text_width)
                    for (wl, _) in wrapped:
                        if is_exhibit_reference(line_str):
                            segments.append({
                                "text": wl,
                                "font_name": "Helvetica-Bold",
                                "font_size": body_font_size,
                                "alignment": "left",
                                "is_heading": False,
                                "is_subheading": False
                            })
                        else:
                            segments.append({
                                "text": wl,
                                "font_name": body_font_name,
                                "font_size": body_font_size,
                                "alignment": "left",
                                "is_heading": False,
                                "is_subheading": False
                            })
            normal_buffer_sec.clear()

        for kind, block_lines in detect_legal_title_blocks(lines_of_body):
            if kind == "legal_page_title_block":
                flush_section_buffer()
                lines_cleaned = [ln.strip() for ln in block_lines]
                segments.append({
                    "legal_page_title": True,
                    "page_always_new": True,
                    "lines": lines_cleaned
                })
            elif kind == "delimiter_line":
                flush_section_buffer()
                segments.append({
                    "delimiter_line": True,
                    "font_name": "Helvetica",
                    "font_size": body_font_size,
                    "is_heading": False,
                    "is_subheading": False
                })
            else:
                normal_buffer_sec.append(block_lines)
        flush_section_buffer()
    return segments

def parse_exhibits_from_text(raw_text):
    lines = raw_text.splitlines()
    pattern = re.compile(r'^\s*EXHIBIT\s+(\d+)\s*:\s*(.*)$', re.IGNORECASE)
    exhibits = OrderedDict()
    current_exhibit_number = None
    current_content = []
    seen_exhibits = set()
    for line in lines:
        match = pattern.match(line)
        if match:
            if current_exhibit_number is not None and current_exhibit_number not in seen_exhibits:
                exhibits[current_exhibit_number] = "\n".join(current_content)
                seen_exhibits.add(current_exhibit_number)
            current_exhibit_number = match.group(1)
            if current_exhibit_number in seen_exhibits:
                current_exhibit_number = None
                current_content = []
                continue
            start_text = match.group(2)
            current_content = [start_text] if start_text else []
        else:
            if current_exhibit_number is not None:
                current_content.append(line)
    if current_exhibit_number is not None and current_exhibit_number not in seen_exhibits:
        exhibits[current_exhibit_number] = "\n".join(current_content)
        seen_exhibits.add(current_exhibit_number)
    return exhibits

def draw_exhibit_text(
    pdf_canvas,
    page_width,
    page_height,
    firm_name,
    case_name,
    exhibit_text,
    exhibit_label,
    page_number,
    total_pages,
    font_name,
    font_size,
    line_spacing
):
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.0 * inch)
    draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
    pdf_canvas.setFont("Helvetica-Bold", 12)
    pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
    pdf_canvas.setLineWidth(1)
    pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
    pdf_canvas.setFont(font_name, font_size)
    max_text_width = page_width - 1.5 * inch
    wrapped = wrap_text_to_lines(pdf_canvas, exhibit_text, font_name, font_size, max_text_width)
    y_text = page_height - 0.8 * inch
    left_margin = 1.0 * inch
    pdf_canvas.setFont("Helvetica-Bold", 10)
    pdf_canvas.drawString(left_margin, y_text, exhibit_label)
    y_text -= (line_spacing * 2)
    pdf_canvas.setFont(font_name, font_size)
    for idx, (txt_line, _) in enumerate(wrapped):
        if y_text < 0.6 * inch:
            pdf_canvas.setFont("Helvetica-Oblique", 9)
            footer_text = f"Page {page_number} of {total_pages}"
            pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)
            pdf_canvas.showPage()
            page_number += 1
            pdf_canvas.setLineWidth(2)
            pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.0 * inch)
            draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
            pdf_canvas.setFont("Helvetica-Bold", 12)
            pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
            pdf_canvas.setLineWidth(1)
            pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
            pdf_canvas.setFont(font_name, font_size)
            y_text = page_height - 0.8 * inch
            pdf_canvas.setFont("Helvetica-Bold", 10)
            pdf_canvas.drawString(left_margin, y_text, exhibit_label)
            y_text -= (line_spacing * 2)
            pdf_canvas.setFont(font_name, font_size)
        pdf_canvas.drawString(left_margin, y_text, txt_line)
        y_text -= line_spacing
    return page_number

def draw_exhibit_image(
    pdf_canvas,
    page_width,
    page_height,
    firm_name,
    case_name,
    image_path,
    page_number,
    total_pages
):
    pdf_canvas.setLineWidth(2)
    pdf_canvas.rect(0.5 * inch, 0.5 * inch, page_width - 1.0 * inch, page_height - 1.0 * inch)
    draw_firm_name_vertical_center(pdf_canvas, firm_name, page_width, page_height)
    pdf_canvas.setFont("Helvetica-Bold", 12)
    pdf_canvas.drawCentredString(page_width / 2.0, page_height - 0.5 * inch, case_name)
    pdf_canvas.setLineWidth(1)
    pdf_canvas.line(0.5 * inch, page_height - 0.6 * inch, page_width - 0.5 * inch, page_height - 0.6 * inch)
    top_of_image_area = page_height - 0.8 * inch
    bottom_of_image_area = 0.5 * inch
    available_height = top_of_image_area - bottom_of_image_area
    available_width = page_width - 1.0 * inch
    if image_path:
        try:
            img_reader = ImageReader(image_path)
            img_width, img_height = img_reader.getSize()
        except Exception as e:
            pdf_canvas.setFont("Helvetica-Oblique", 10)
            pdf_canvas.drawCentredString(
                page_width / 2.0,
                page_height / 2.0,
                f"Unable to load image: {image_path} Error: {e}"
            )
        else:
            scale = min(available_width / img_width, available_height / img_height, 1.0)
            new_width = img_width * scale
            new_height = img_height * scale
            x_img = 0.5 * inch + (available_width - new_width) / 2.0
            y_img_bottom = bottom_of_image_area
            pdf_canvas.drawImage(
                img_reader,
                x_img,
                y_img_bottom,
                width=new_width,
                height=new_height,
                preserveAspectRatio=True,
                anchor='c'
            )
    pdf_canvas.setFont("Helvetica-Oblique", 9)
    footer_text = f"Page {page_number} of {total_pages}"
    pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)

class Lawsuit:
    def __init__(
        self,
        sections=None,
        exhibits=None,
        header=None,
        documents=None,
        case_information="",
        law_firm_information=""
    ):
        if sections is None:
            sections = OrderedDict()
        if exhibits is None:
            exhibits = OrderedDict()
        if header is None:
            header = OrderedDict()
        if documents is None:
            documents = OrderedDict()
        self.sections = OrderedDict(sections)
        self.exhibits = OrderedDict(exhibits)
        self.header = OrderedDict(header)
        self.documents = OrderedDict(documents)
        self.case_information = case_information
        self.law_firm_information = law_firm_information
        self.tf_model = self._build_tf_model()
        self._agi_model = self._build_agi_model()
        self.ai_legal_notes = ""
        self.agi_legal_professional_output = ""
        self.law_firm_information = self._transform_text_with_tensorflow(self.law_firm_information)

    def _build_tf_model(self):
        model = tf.keras.Sequential([
            layers.Input(shape=(128,)),
            layers.Dense(64, activation='relu'),
            layers.Dense(1, activation='sigmoid')
        ])
        model.compile(optimizer='adam', loss='binary_crossentropy', metrics=['accuracy'])
        return model

    def _build_agi_model(self):
        model = tf.keras.Sequential([
            layers.Input(shape=(256,)),
            layers.Dense(128, activation='relu'),
            layers.Dense(64, activation='relu'),
            layers.Dense(1, activation='sigmoid')
        ])
        model.compile(optimizer='adam', loss='binary_crossentropy', metrics=['accuracy'])
        return model

    def _transform_text_with_tensorflow(self, text):
        return f"EnhancedTF({text})"

    def run_deep_legal_analysis(self):
        combined_text = []
        for sec_value in self.sections.values():
            combined_text.append(sec_value)
        raw_input_data = " ".join(combined_text)
        self.ai_legal_notes = "AGI analysis: " + raw_input_data[:50] + "..."
        dummy_vector = tf.zeros((1, 128))
        _ = self.tf_model.predict(dummy_vector)

    def run_agi_legal_professionalism(self, pdf_file_texts):
        aggregated_text = " ".join(pdf_file_texts)[:100]
        self.agi_legal_professional_output = "Advanced AGI reply: " + aggregated_text
        dummy_input = tf.zeros((1, 256))
        _ = self._agi_model.predict(dummy_input)

    def __repr__(self):
        header_str = "\n".join([f"  {k}: {v}" for k, v in self.header.items()])
        sections_str = "\n".join([f"  {sec_key}: {sec_value}" for sec_key, sec_value in self.sections.items()])
        exhibits_str = []
        for ex_key, ex_data in self.exhibits.items():
            ex_inner = "\n      ".join([f"{ik}: {iv}" for ik, iv in ex_data.items()])
            exhibits_str.append(f"  {ex_key}:\n      {ex_inner}")
        exhibits_str = "\n".join(exhibits_str)
        documents_str = []
        for doc_key, doc_text in self.documents.items():
            documents_str.append(f"  {doc_key}:\n      {doc_text}")
        documents_str = "\n".join(documents_str)
        return (
            "Lawsuit Object:\n\n"
            "CASE INFORMATION:\n"
            f"  {self.case_information}\n\n"
            "LAW FIRM INFORMATION:\n"
            f"  {self.law_firm_information}\n\n"
            "HEADER:\n"
            f"{header_str}\n\n"
            "SECTIONS:\n"
            f"{sections_str}\n\n"
            "EXHIBITS:\n"
            f"{exhibits_str}\n\n"
            "DOCUMENTS:\n"
            f"{documents_str}\n\n"
            "AI LEGAL NOTES:\n"
            f"  {self.ai_legal_notes}\n\n"
            "AGI LEGAL PROFESSIONALISM:\n"
            f"  {self.agi_legal_professional_output}\n"
        )

def detect_case_numbers(text):
    pattern = re.compile(r'\b([A-Z]{1,5}\s*\d{1,}-\d+)\b', re.IGNORECASE)
    return set(re.findall(pattern, text))

def store_lawsuit_in_db(lawsuit_obj, db_conn):
    db_conn.execute("""
        CREATE TABLE IF NOT EXISTS cases (
            case_number TEXT PRIMARY KEY,
            firm_name TEXT,
            creation_date TEXT,
            data BLOB,
            is_active INTEGER DEFAULT 0
        )
    """)
    pickled_data = pickle.dumps(lawsuit_obj)
    existing = db_conn.execute(
        "SELECT case_number FROM cases WHERE case_number = ?",
        (lawsuit_obj.case_information,)
    ).fetchone()
    if existing:
        db_conn.execute(
            "UPDATE cases SET firm_name=?, creation_date=?, data=? WHERE case_number=?",
            (lawsuit_obj.law_firm_information, datetime.datetime.now().isoformat(), pickled_data, lawsuit_obj.case_information)
        )
    else:
        db_conn.execute(
            "INSERT INTO cases (case_number, firm_name, creation_date, data) VALUES (?, ?, ?, ?)",
            (lawsuit_obj.case_information, lawsuit_obj.law_firm_information, datetime.datetime.now().isoformat(), pickled_data)
        )
    db_conn.commit()

def store_detected_cases_in_db(detected_cases, db_conn):
    db_conn.execute("""
        CREATE TABLE IF NOT EXISTS detected_cases (
            case_number TEXT PRIMARY KEY,
            detection_date TEXT
        )
    """)
    for cn in detected_cases:
        existing = db_conn.execute(
            "SELECT case_number FROM detected_cases WHERE case_number = ?",
            (cn,)
        ).fetchone()
        if not existing:
            db_conn.execute(
                "INSERT INTO detected_cases (case_number, detection_date) VALUES (?, ?)",
                (cn, datetime.datetime.now().isoformat())
            )
    db_conn.commit()

def set_active_case(case_number, db_conn):
    db_conn.execute("CREATE TABLE IF NOT EXISTS cases (case_number TEXT PRIMARY KEY, firm_name TEXT, creation_date TEXT, data BLOB, is_active INTEGER DEFAULT 0)")
    db_conn.execute("UPDATE cases SET is_active = 0")
    db_conn.execute("UPDATE cases SET is_active = 1 WHERE case_number = ?", (case_number,))
    db_conn.commit()

def generate_legal_document(
    firm_name,
    case_name,
    output_filename,
    header_od,
    sections_od,
    exhibits,
    heading_positions
):
    page_width, page_height = letter
    pdf_canvas = canvas.Canvas(output_filename, pagesize=letter)
    pdf_canvas.setTitle("Legal Document without Cover Sheet")
    pdf_canvas.setAuthor(firm_name)
    pdf_canvas.setSubject(case_name)
    pdf_canvas.setCreator("Legal PDF Generator")

    heading_styles = classify_headings(sections_od)
    top_margin = 1.0 * inch
    bottom_margin = 1.0 * inch
    left_margin = 1.2 * inch
    right_margin = 0.5 * inch
    line_spacing = 0.25 * inch
    usable_height = page_height - (top_margin + bottom_margin)
    max_lines_per_page = int(usable_height // line_spacing)
    line_offset_x = left_margin
    line_offset_y = page_height - top_margin
    max_text_width = page_width - right_margin - line_offset_x - 0.2 * inch

    segments = prepare_main_pdf_segments(
        header_text=header_od.get("content", ""),
        sections_od=sections_od,
        heading_styles=heading_styles,
        pdf_canvas=pdf_canvas,
        max_text_width=max_text_width
    )
    current_index = 0
    text_pages = 0
    total_segments = len(segments)
    while current_index < total_segments:
        seg = segments[current_index]
        if seg.get("page_always_new"):
            text_pages += 1
            current_index += 1
        else:
            lines_used = 0
            local_i = current_index
            while local_i < total_segments and lines_used < max_lines_per_page:
                s = segments[local_i]
                if s.get("page_always_new"):
                    break
                lines_used += 1
                local_i += 1
            text_pages += 1
            current_index = local_i
    exhibit_pages_est = 0
    if exhibits:
        exhibit_pages_est = len(exhibits)
    total_pages_est = text_pages + exhibit_pages_est * 2
    page_number = 1
    current_index = 0
    while current_index < total_segments:
        next_index = draw_page_of_segments(
            pdf_canvas=pdf_canvas,
            page_width=page_width,
            page_height=page_height,
            segments=segments,
            start_index=current_index,
            max_lines_per_page=max_lines_per_page,
            firm_name=firm_name,
            case_name=case_name,
            page_number=page_number,
            total_pages=total_pages_est,
            line_offset_x=line_offset_x,
            line_offset_y=line_offset_y,
            line_spacing=line_spacing,
            heading_positions=heading_positions
        )
        pdf_canvas.showPage()
        page_number += 1
        current_index = next_index

    idx = 0
    for ex_content, image_path in exhibits:
        idx += 1
        exhibit_label = f"EXHIBIT {idx}:"
        page_number = draw_exhibit_text(
            pdf_canvas=pdf_canvas,
            page_width=page_width,
            page_height=page_height,
            firm_name=firm_name,
            case_name=case_name,
            exhibit_text=ex_content,
            exhibit_label=exhibit_label,
            page_number=page_number,
            total_pages=total_pages_est,
            font_name="Helvetica",
            font_size=10,
            line_spacing=line_spacing
        )
        pdf_canvas.setFont("Helvetica-Oblique", 9)
        footer_text = f"Page {page_number} of {total_pages_est}"
        pdf_canvas.drawCentredString(page_width / 2.0, 0.4 * inch, footer_text)
        pdf_canvas.showPage()
        page_number += 1
        if image_path:
            draw_exhibit_image(
                pdf_canvas=pdf_canvas,
                page_width=page_width,
                page_height=page_height,
                firm_name=firm_name,
                case_name=case_name,
                image_path=image_path,
                page_number=page_number,
                total_pages=total_pages_est
            )
            pdf_canvas.showPage()
            page_number += 1

    pdf_canvas.save()
    generate_complaint_docx(
        docx_filename=os.path.splitext(output_filename)[0] + ".docx",
        firm_name=firm_name,
        case_name=case_name,
        header_od=header_od,
        sections_od=sections_od,
        heading_styles=heading_styles
    )

def filter_headings_for_toc(heading_positions):
    new_positions = []
    found_exhibit = False
    in_special_section = False
    for heading_text, pg_num, ln_num, is_sub in heading_positions:
        if in_special_section:
            new_positions.append((heading_text, pg_num, ln_num, is_sub))
            continue
        if found_exhibit:
            if re.match(r'(?i)^SPECIAL EXHIBITS$', heading_text.strip()):
                in_special_section = True
                new_positions.append((heading_text, pg_num, ln_num, is_sub))
            continue
        new_positions.append((heading_text, pg_num, ln_num, is_sub))
        if is_exhibit_reference(heading_text):
            found_exhibit = True
    return new_positions

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--firm_name", default="PDFSage Inc.")
    parser.add_argument("--case", required=True)
    parser.add_argument("--output", default="lawsuit.pdf")
    parser.add_argument("--file", required=True)
    parser.add_argument("--index", default="index.pdf")
    parser.add_argument("--pickle", nargs='?', const=None)
    parser.add_argument("--set-case", help="Set the specified case number as active in the database", required=False)
    parser.add_argument("--reply", nargs='*', help="Reply with advanced tensorflow if PDF or ZIP is provided")
    parser.add_argument("--exhibits", nargs='*', help="Optional image paths for exhibits")
    args = parser.parse_args()

    datetime_string = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    raw_text = read_input_file(args.file)

    detected_cases = detect_case_numbers(raw_text)
    db_conn = sqlite3.connect("cases.db")
    store_detected_cases_in_db(detected_cases, db_conn)

    def separate_after_exhibit_1(full_text):
        pattern_ex1 = re.compile(r'^\s*EXHIBIT\s+1\s*:', re.IGNORECASE)
        lines = full_text.splitlines()
        main_part = []
        exhibits_part = []
        found_ex1 = False
        for line in lines:
            if not found_ex1 and pattern_ex1.match(line):
                found_ex1 = True
                exhibits_part.append(line)
            elif found_ex1:
                exhibits_part.append(line)
            else:
                main_part.append(line)
        return "\n".join(main_part), "\n".join(exhibits_part)

    main_text, exhibit_text_after_1 = separate_after_exhibit_1(raw_text)
    text_exhibits_od = parse_exhibits_from_text(exhibit_text_after_1)

    header_od, sections_od = parse_header_and_sections(main_text)

    exhibits_od = OrderedDict()
    i = 1
    for ex_key in sorted(text_exhibits_od.keys(), key=lambda x: int(x)):
        exhibits_od[str(i)] = OrderedDict([
            ('caption', text_exhibits_od[ex_key]),
            ('image_path', "")
        ])
        i += 1

    header_od["DocumentTitle"] = "Complaint for Tort – Other"
    header_od["DateFiled"] = "2025-02-14"
    header_od["Court"] = "King County Superior Court"

    found_documents = parse_documents_from_text(raw_text)
    documents_od = OrderedDict()
    for idx, doc_text in enumerate(found_documents, start=1):
        documents_od[str(idx)] = doc_text

    args.output = generate_smart_filename(args.output, raw_text, datetime_string)
    args.index = generate_smart_filename(args.index, raw_text, datetime_string)
    if args.pickle is not None:
        if args.pickle:
            args.pickle = generate_smart_filename(args.pickle, raw_text, datetime_string)
        else:
            default_pickle = f"lawsuit.pickle"
            args.pickle = generate_smart_filename(default_pickle, raw_text, datetime_string)

    # Create the Lawsuit object
    lawsuit_obj = Lawsuit(
        sections=sections_od,
        exhibits=exhibits_od,
        header=header_od,
        documents=documents_od,
        case_information=args.case,
        law_firm_information=args.firm_name
    )

    # Update exhibit image paths if provided
    if args.exhibits:
        i = 1
        for ex_image in args.exhibits:
            ex_key = str(i)
            if ex_key in lawsuit_obj.exhibits:
                lawsuit_obj.exhibits[ex_key]['image_path'] = ex_image
            i += 1

    store_lawsuit_in_db(lawsuit_obj, db_conn)
    if args.set_case:
        set_active_case(args.set_case, db_conn)

    lawsuit_obj.run_deep_legal_analysis()

    if args.reply:
        pdf_file_texts = []
        for reply_file in args.reply:
            ext = os.path.splitext(reply_file)[1].lower()
            if os.path.isfile(reply_file) and (ext == '.pdf' or ext == '.zip'):
                if ext == '.pdf':
                    pdf_file_texts.append(read_input_file(reply_file))
                if ext == '.zip':
                    import zipfile
                    with zipfile.ZipFile(reply_file, 'r') as z:
                        for name in z.namelist():
                            if name.lower().endswith('.pdf'):
                                import io
                                from PyPDF2 import PdfReader
                                with z.open(name) as fpdf:
                                    pdf_bytes = fpdf.read()
                                    pdf_stream = io.BytesIO(pdf_bytes)
                                    reader = PdfReader(pdf_stream)
                                    lines = []
                                    for page in reader.pages:
                                        text = page.extract_text()
                                        if text:
                                            lines.extend(text.splitlines())
                                    pdf_file_texts.append("\n".join(lines))
        if pdf_file_texts:
            lawsuit_obj.run_agi_legal_professionalism(pdf_file_texts)

    exhibits_for_pdf = []
    for _, val in lawsuit_obj.exhibits.items():
        exhibits_for_pdf.append((val["caption"], val["image_path"]))

    heading_positions = []
    generate_legal_document(
        firm_name=args.firm_name,
        case_name=args.case,
        output_filename=args.output,
        header_od=header_od,
        sections_od=sections_od,
        exhibits=exhibits_for_pdf,
        heading_positions=heading_positions
    )

    heading_positions = filter_headings_for_toc(heading_positions)
    generate_index_pdf(
        index_filename=args.index,
        firm_name=args.firm_name,
        case_name=args.case,
        heading_positions=heading_positions
    )
    index_docx = os.path.splitext(args.index)[0] + ".docx"
    generate_toc_docx(
        docx_filename=index_docx,
        firm_name=args.firm_name,
        case_name=args.case,
        heading_positions=heading_positions
    )

    if args.pickle is not None:
        with open(args.pickle, "wb") as pf:
            pickle.dump(lawsuit_obj, pf)
        pkl_path = args.pickle
    else:
        pkl_path = "Not saved (not requested)."

    print(f"PDF generated (exhibits + pre-exhibit content): {args.output}")
    print(f"DOCX Complaint generated: {os.path.splitext(args.output)[0] + '.docx'}")
    print(f"Index PDF generated: {args.index}")
    print(f"Index DOCX generated: {index_docx}")
    print(f"Lawsuit object saved to: {pkl_path}\n")
    print("Dumped Lawsuit object:")
    print(lawsuit_obj)

    db_conn.close()

if __name__ == "__main__":
    main()